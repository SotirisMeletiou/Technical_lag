import docx
from packaging import version
import pandas as pd
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL

def read_file(file_path):
    with open(file_path, 'r') as file:
        lines = file.readlines()
    return lines

def parse_data(lines):
    data = [line.strip().split() for line in lines]
    return sorted(data, key=lambda x: version.parse(x[0]))

def create_table(document, data):
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Add table header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Version'
    header_cells[1].text = 'Value1'
    header_cells[2].text = 'Value2'

    # Add table rows
    for row_data in data:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]

def add_dataframe_as_table(doc, df, title):
    doc.add_paragraph(title)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = 'Table Grid'

    # Add column headers
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = col
        cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)

    # Add data
    for i, value in enumerate(df.values):
        for j, val in enumerate(value):
            cell = table.cell(i + 1, j)
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)

    return table

if __name__ == "__main__":
    file_path_evolve = "evolve.txt"
    file_path_max_lag = "max_lag.txt"
    output_path = "output_combined.docx"

    # Read and process evolve.txt
    lines_evolve = read_file(file_path_evolve)
    sorted_data_evolve = parse_data(lines_evolve)

    # Read and process max_lag.txt
    data_max_lag = pd.read_csv(file_path_max_lag, header=None, names=['version', 'dependency', 'lag'])

    # Create a Word document
    document = docx.Document()

    # Create table for evolve.txt
    create_table(document, sorted_data_evolve)

    # Add a paragraph to separate the tables
    document.add_paragraph()

    # Add the DataFrame from max_lag.txt as a table
    add_dataframe_as_table(document, data_max_lag, 'Max Lag Table')

    # Save the combined document
    document.save(output_path)
    print(f"Combined document has been created and saved to {output_path}")
